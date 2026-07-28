import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

import zstandard
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from converted_document import DocumentConversionError
from converters import convert_document, get_converter
from converters.base import ConversionLimits
from converters.docx import DocxConverter
from converters.pdf import PdfConverter
from source_map import strip_source_unit_markers


def read_source_map(path: Path) -> list[dict]:
    with (
        path.open("rb") as raw,
        zstandard.ZstdDecompressor().stream_reader(raw) as reader,
    ):
        payload = reader.read().decode("utf-8")
    return [json.loads(line) for line in payload.splitlines()]


def add_pdf_text_page(writer: PdfWriter, text: str, *, rotate: bool = False) -> None:
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_reference})}
    )
    stream = DecodedStreamObject()
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream.set_data(f"BT /F1 18 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(stream)
    if rotate:
        page.rotate(90)


def write_pdf(path: Path, texts: list[str | None], *, encrypted: bool = False) -> None:
    writer = PdfWriter()
    for text in texts:
        if text is None:
            writer.add_blank_page(width=612, height=792)
        else:
            add_pdf_text_page(writer, text)
    if encrypted:
        writer.encrypt("local-test-password")
    with path.open("wb") as stream:
        writer.write(stream)


def copy_zip_with_extra(source: Path, target: Path, name: str, payload: bytes) -> None:
    with zipfile.ZipFile(source, "r") as existing, zipfile.ZipFile(
        target, "w", compression=zipfile.ZIP_DEFLATED
    ) as output:
        for entry in existing.infolist():
            output.writestr(entry, existing.read(entry))
        output.writestr(name, payload)


class PdfConverterTests(unittest.TestCase):
    def setUp(self):
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary_directory.name)

    def tearDown(self):
        self.temporary_directory.cleanup()

    def test_text_pdf_is_deterministic_and_retains_page_locator(self):
        source = self.root / "manual.pdf"
        write_pdf(source, ["Hello from page one", None])

        first = convert_document(source, self.root / "first")
        second = PdfConverter().convert(source, self.root / "second")

        markdown = first.document.path.read_text(encoding="utf-8")
        self.assertEqual(strip_source_unit_markers(markdown), "Hello from page one\n")
        self.assertEqual(first.document.sha256, second.document.sha256)
        self.assertEqual(first.source_map.sha256, second.source_map.sha256)
        self.assertEqual(first.manifest_artifact.sha256, second.manifest_artifact.sha256)
        self.assertEqual(first.manifest.document_kind, "pdf")
        self.assertEqual(first.manifest.conversion_profile, "pdf-text-v1")
        self.assertEqual(first.manifest.source_encoding, "binary")
        self.assertEqual(
            first.manifest.warnings,
            ("PDF_SOME_PAGES_HAVE_NO_EXTRACTABLE_TEXT",),
        )

        units = read_source_map(first.source_map.path)
        self.assertEqual(len(units), 1)
        self.assertEqual(
            units[0]["source"],
            {"block": 1, "kind": "page_text", "page": 1, "type": "pdf"},
        )
        byte_range = units[0]["markdown"]
        extracted = first.document.path.read_bytes()[
            byte_range["byte_start"] : byte_range["byte_end"]
        ].decode("utf-8")
        self.assertEqual(extracted, "Hello from page one\n")

    def test_rotated_text_page_emits_complex_layout_warning(self):
        source = self.root / "rotated.pdf"
        writer = PdfWriter()
        add_pdf_text_page(writer, "Rotated but extractable", rotate=True)
        with source.open("wb") as stream:
            writer.write(stream)

        result = convert_document(source, self.root / "derived")

        self.assertIn("Rotated but extractable", result.document.path.read_text(encoding="utf-8"))
        self.assertEqual(result.manifest.warnings, ("PDF_COMPLEX_LAYOUT_MAY_BE_LOSSY",))

    def test_pdf_rejects_invalid_signature_encryption_and_missing_text_layer(self):
        invalid = self.root / "invalid.pdf"
        invalid.write_bytes(b"not a PDF")

        encrypted = self.root / "encrypted.pdf"
        write_pdf(encrypted, ["secret"], encrypted=True)

        blank = self.root / "blank.pdf"
        write_pdf(blank, [None])

        cases = (
            (invalid, "PDF_INVALID_SIGNATURE"),
            (encrypted, "PDF_ENCRYPTED"),
            (blank, "PDF_HAS_NO_EXTRACTABLE_TEXT"),
        )
        for index, (source, expected_code) in enumerate(cases):
            output = self.root / f"rejected-{index}"
            with self.subTest(expected_code=expected_code):
                with self.assertRaises(DocumentConversionError) as context:
                    convert_document(source, output)
                self.assertEqual(context.exception.code, expected_code)
                self.assertFalse((output / "document.md").exists())

    def test_pdf_enforces_page_and_extracted_text_budgets(self):
        too_many_pages = self.root / "too-many-pages.pdf"
        write_pdf(too_many_pages, ["one", "two"])
        with self.assertRaises(DocumentConversionError) as page_context:
            PdfConverter(
                ConversionLimits(max_pdf_pages=1),
            ).convert(too_many_pages, self.root / "page-limit")
        self.assertEqual(page_context.exception.code, "PDF_TOO_MANY_PAGES")

        too_much_text = self.root / "too-much-text.pdf"
        write_pdf(too_much_text, ["bounded extraction text"])
        with self.assertRaises(DocumentConversionError) as text_context:
            PdfConverter(
                ConversionLimits(max_pdf_extracted_chars=5),
            ).convert(too_much_text, self.root / "text-limit")
        self.assertEqual(text_context.exception.code, "PDF_TEXT_TOO_LARGE")


class DocxConverterTests(unittest.TestCase):
    def setUp(self):
        self.temporary_directory = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary_directory.name)

    def tearDown(self):
        self.temporary_directory.cleanup()

    def _write_ordered_document(self, path: Path) -> None:
        document = Document()
        document.add_heading("Operations Guide", level=1)
        document.add_paragraph("First paragraph.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Key"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Retries"
        table.cell(1, 1).text = "3 | attempts"
        document.add_paragraph("Final paragraph.")
        document.save(path)

    def test_docx_preserves_body_order_and_paragraph_table_locators(self):
        source = self.root / "guide.docx"
        self._write_ordered_document(source)

        first = convert_document(source, self.root / "first")
        second = DocxConverter().convert(source, self.root / "second")

        markdown = strip_source_unit_markers(first.document.path.read_text(encoding="utf-8"))
        expected_fragments = (
            "# Operations Guide",
            "First paragraph.",
            "| Key | Value |",
            "| Retries | 3 \\| attempts |",
            "Final paragraph.",
        )
        positions = [markdown.index(fragment) for fragment in expected_fragments]
        self.assertEqual(positions, sorted(positions))
        self.assertEqual(first.document.sha256, second.document.sha256)
        self.assertEqual(first.source_map.sha256, second.source_map.sha256)
        self.assertEqual(first.manifest_artifact.sha256, second.manifest_artifact.sha256)
        self.assertEqual(first.manifest.document_kind, "docx")
        self.assertEqual(first.manifest.conversion_profile, "docx-v1")

        sources = [unit["source"] for unit in read_source_map(first.source_map.path)]
        self.assertEqual(sources[0], {"heading_level": 1, "kind": "heading", "paragraph": 1, "type": "docx"})
        self.assertEqual(sources[1], {"kind": "paragraph", "paragraph": 2, "type": "docx"})
        self.assertEqual(
            sources[2],
            {
                "column_count": 2,
                "kind": "table",
                "row_end": 2,
                "row_start": 1,
                "table": 1,
                "type": "docx",
            },
        )
        self.assertEqual(sources[3], {"kind": "paragraph", "paragraph": 3, "type": "docx"})

    def test_docx_images_are_ignored_without_ocr_and_reported(self):
        source = self.root / "base.docx"
        Document().save(source)
        with_text = Document()
        with_text.add_paragraph("Visible document text")
        with_text.save(source)
        image_container = self.root / "image.docx"
        copy_zip_with_extra(source, image_container, "word/media/image1.png", b"TEXT IN IMAGE")

        result = convert_document(image_container, self.root / "derived")

        markdown = result.document.path.read_text(encoding="utf-8")
        self.assertIn("Visible document text", markdown)
        self.assertNotIn("TEXT IN IMAGE", markdown)
        self.assertEqual(result.manifest.warnings, ("DOCX_IMAGES_IGNORED",))

    def test_docx_rejects_invalid_ooxml_macros_unsafe_paths_and_zip_bombs(self):
        valid = self.root / "valid.docx"
        self._write_ordered_document(valid)

        invalid_signature = self.root / "invalid.docx"
        invalid_signature.write_bytes(b"not a ZIP")

        missing_parts = self.root / "missing.docx"
        with zipfile.ZipFile(missing_parts, "w") as archive:
            archive.writestr("word/document.xml", b"<document/>")

        macro = self.root / "macro.docx"
        copy_zip_with_extra(valid, macro, "word/vbaProject.bin", b"macro")

        unsafe_path = self.root / "unsafe.docx"
        with zipfile.ZipFile(unsafe_path, "w") as archive:
            archive.writestr("../escape.xml", b"<xml/>")

        zip_bomb = self.root / "bomb.docx"
        with zipfile.ZipFile(zip_bomb, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", b"A" * (1024 * 1024))

        cases = (
            (invalid_signature, "DOCX_INVALID_SIGNATURE"),
            (missing_parts, "DOCX_INVALID_OOXML"),
            (macro, "DOCX_MACROS_NOT_ALLOWED"),
            (unsafe_path, "DOCX_UNSAFE_ZIP_ENTRY"),
            (zip_bomb, "DOCX_ZIP_COMPRESSION_RATIO_EXCEEDED"),
        )
        for index, (source, expected_code) in enumerate(cases):
            output = self.root / f"rejected-{index}"
            with self.subTest(expected_code=expected_code):
                with self.assertRaises(DocumentConversionError) as context:
                    convert_document(source, output)
                self.assertEqual(context.exception.code, expected_code)
                self.assertFalse((output / "document.md").exists())

    def test_empty_docx_has_stable_empty_document_error(self):
        source = self.root / "empty.docx"
        Document().save(source)

        with self.assertRaises(DocumentConversionError) as context:
            convert_document(source, self.root / "derived")

        self.assertEqual(context.exception.code, "EMPTY_DOCUMENT")


class RegistryTests(unittest.TestCase):
    def test_registry_resolves_pdf_and_docx_case_insensitively(self):
        self.assertIsInstance(get_converter("FILE.PDF"), PdfConverter)
        self.assertIsInstance(get_converter("FILE.DOCX"), DocxConverter)


if __name__ == "__main__":
    unittest.main()
